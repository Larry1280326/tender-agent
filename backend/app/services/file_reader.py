"""本地檔案文字抽取：將上傳／下載嘅文件（pdf/docx/xlsx/doc/xls/txt/csv）轉做純文字。

每個 parser 都係 lazy import，咁就算缺某一隻 lib 都唔會搞冧成個 module；
caller（read_file tool）自己 catch ValueError／ImportError 轉做清晰錯誤。
"""
from __future__ import annotations

import re
from pathlib import Path

# 可抽取文字嘅副檔名（read_file 用嚟校驗；download 白名單仲有 zip/ppt/rar 等，但唔係文字）。
READABLE_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".doc", ".xls", ".txt", ".csv", ".md",
}

# .doc（OLE2）文字抽取：UTF-16LE 解碼後，捉「文字似」嘅連續 run（中英數字＋標點＋空白）。
_DOC_RUN_RE = re.compile(
    r"[一-鿿　-〿぀-ヿA-Za-z0-9\s.,;:!?()\[\]{}<>'\"&%$#@/\\_\-=+|~^`‐-’“-”]{4,}"
)
_DOC_ASCII_RE = re.compile(r"[A-Za-z0-9\s.,;:!?()\[\]{}<>'\"&%$#@/\\_\-=+|~^`]{4,}")


def _pdf_text(path: Path) -> str:
    try:
        import pymupdf
    except ImportError:  # 舊版叫 fitz
        import fitz as pymupdf  # type: ignore[no-redef]
    doc = pymupdf.open(str(path))
    try:
        return "\n\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _docx_text(path: Path) -> str:
    from docx import Document  # python-docx

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    out: list[str] = []
    try:
        for ws in wb.worksheets:
            out.append(f"# {ws.title}")
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if vals:
                    out.append(" | ".join(vals))
    finally:
        wb.close()
    return "\n".join(out)


def _xls_text(path: Path) -> str:
    import xlrd  # 1.2.0（最後支援 .xls 嘅版本）

    wb = xlrd.open_workbook(str(path))
    out: list[str] = []
    for sh in wb.sheets():
        out.append(f"# {sh.name}")
        for r in range(sh.nrows):
            vals = [str(sh.cell_value(r, c)).strip() for c in range(sh.ncols)]
            vals = [v for v in vals if v]
            if vals:
                out.append(" | ".join(vals))
    return "\n".join(out)


def _doc_runs(data: bytes) -> str:
    """由 .doc 嘅 WordDocument stream 抽文字 run（best-effort，複雜版式可能唔完整）。"""
    chunks: list[str] = []
    try:
        text16 = data.decode("utf-16-le", errors="ignore")
        chunks.extend(_DOC_RUN_RE.findall(text16))
    except Exception:  # noqa: BLE001
        pass
    if not chunks:
        text8 = data.decode("cp1252", errors="ignore")
        chunks.extend(_DOC_ASCII_RE.findall(text8))
    return "\n".join(chunks)


def _doc_text(path: Path) -> str:
    import olefile

    ole = olefile.OleFileIO(str(path))
    try:
        data = b""
        for stream in ("WordDocument", "1Table", "0Table"):
            if ole.exists(stream):
                data = ole.openstream(stream).read()
                if data:
                    break
    finally:
        ole.close()
    text = _doc_runs(data)
    if not text.strip():
        return "（.doc 抽取唔到文字；建議轉 .docx 或 .pdf 再上傳以取得完整內容）"
    return text


def _text_file(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "gb18030", "big5"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return raw.decode("latin-1", errors="replace")


_DISPATCH = {
    ".pdf": _pdf_text,
    ".docx": _docx_text,
    ".xlsx": _xlsx_text,
    ".xls": _xls_text,
    ".doc": _doc_text,
    ".txt": _text_file,
    ".csv": _text_file,
    ".md": _text_file,
}


def extract_text(path: Path) -> str:
    """按副檔名抽取文字；唔支援嘅副檔名／缺 lib 就 raise ValueError（caller 轉做清晰訊息）。"""
    ext = path.suffix.lower()
    fn = _DISPATCH.get(ext)
    if fn is None:
        supported = " / ".join(sorted(READABLE_EXTENSIONS))
        raise ValueError(f"不支援嘅副檔名：{ext or '(無)'}（只支援 {supported}）。")
    try:
        return fn(path)
    except ImportError as e:
        raise ValueError(f"缺少解析 {ext} 所需嘅函式庫：{e}") from e
